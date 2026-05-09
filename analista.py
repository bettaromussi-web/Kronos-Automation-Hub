import sys
import os
import pandas as pd
import numpy as np
import torch
import plotly.graph_objects as go
import requests
from docx import Document
from docx.shared import Inches
from datetime import datetime

# --- CONFIGURAZIONE TELEGRAM ---
def send_telegram_alert(message):
    token = os.getenv('TELEGRAM_TOKEN')
    chat_id = os.getenv('TELEGRAM_CHAT_ID')
    if not token or not chat_id:
        print("⚠️ Telegram non configurato (Secret mancanti).")
        return
    url = f"https://telegram.org{token}/sendMessage"
    payload = {"chat_id": chat_id, "text": message, "parse_mode": "Markdown"}
    try:
        requests.post(url, data=payload)
    except Exception as e:
        print(f"❌ Errore invio Telegram: {e}")

# --- CONFIGURAZIONE SISTEMA ---
import argparse
parser = argparse.ArgumentParser()
parser.add_argument('--sector', type=str, required=True)
args = parser.parse_args()

SETTORE = args.sector.lower()
CARTELLA_DATI = f"database_{SETTORE}"
REPORT_DIR = "reports"
LOOKBACK = 512
PRED_LEN = 120
DEVICE = "cpu" # GitHub Actions usa CPU gratuitamente

if not os.path.exists(REPORT_DIR): os.makedirs(REPORT_DIR)

# Caricamento Kronos
sys.path.append(os.getcwd())
from model import Kronos, KronosTokenizer, KronosPredictor
tokenizer = KronosTokenizer.from_pretrained("NeoQuasar/Kronos-Tokenizer-base")
model = Kronos.from_pretrained("NeoQuasar/Kronos-small").to(DEVICE)
predictor = KronosPredictor(model, tokenizer, device=DEVICE, max_context=LOOKBACK)

# Creazione Documento Word
doc = Document()
doc.add_heading(f'Report Kronos: Settore {SETTORE.upper()}', 0)

titoli_cartella = [f.replace('.csv', '') for f in os.listdir(CARTELLA_DATI) if f.endswith('.csv')]
classifica_risultati = []

for ticker in titoli_cartella:
    print(f"🔍 Analisi {ticker}...")
    df = pd.read_csv(os.path.join(CARTELLA_DATI, f"{ticker}.csv"))
    
    x_df = df.iloc[-LOOKBACK:].copy()
    x_df.columns = [c.lower() for c in x_df.columns]
    x_ts = pd.to_datetime(df['Datetime'].iloc[-LOOKBACK:])
    y_ts = pd.Series(pd.date_range(start=x_ts.iloc[-1], periods=PRED_LEN + 1, freq='5min')[1:])

    try:
        pred = predictor.predict(x_df[['open','high','low','close','volume']], x_ts, y_ts, PRED_LEN, T=0.8, sample_count=5)
        
        var_perc = ((pred['close'].iloc[-1] - x_df['close'].iloc[-1]) / x_df['close'].iloc[-1]) * 100
        sentimento = "🚀 RIALZISTA" if var_perc > 0.5 else "🔻 RIBASSISTA" if var_perc < -0.5 else "➡️ STABILE"
        classifica_risultati.append({"ticker": ticker, "var": var_perc, "sent": sentimento})

        # Grafico per Word
        fig = go.Figure()
        fig.add_trace(go.Scatter(y=x_df['close'].tail(50), name='Storico'))
        fig.add_trace(go.Scatter(x=list(range(50, 50+PRED_LEN)), y=pred['close'], name='Predizione', line=dict(dash='dash', color='red')))
        img_path = f"temp_{ticker}.png"
        fig.write_image(img_path)

        doc.add_heading(f'Titolo: {ticker}', level=1)
        doc.add_picture(img_path, width=Inches(5))
        doc.add_paragraph(f"Sentimento: {sentimento} | Variazione prevista: {var_perc:.2f}%")
        os.remove(img_path)
    except Exception as e:
        print(f"⚠️ Errore su {ticker}: {e}")

# Salva Report
nome_file = os.path.join(REPORT_DIR, f"Report_{SETTORE.upper()}.docx")
doc.save(nome_file)

# --- INVIO ALERT TELEGRAM ---
if classifica_risultati:
    # Ordina per rendimento decrescente
    classifica_risultati.sort(key=lambda x: x['var'], reverse=True)
    
    messaggio = f"📊 *REPORT KRONOS: {SETTORE.upper()}*\n\n"
    for r in classifica_risultati:
        messaggio += f"{r['sent']} *{r['ticker']}*: {r['var']:.2f}%\n"
    
    messaggio += f"\n📂 Report completo generato correttamente."
    send_telegram_alert(messaggio)

print(f"\n✅ Analisi completata per {SETTORE}.")
